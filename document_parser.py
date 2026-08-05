"""
Document parser for ATS Bot
Supports PDF, DOCX, and TXT files and extracts embedded images when possible.
"""

import io
import os
import PyPDF2
from docx import Document
from pathlib import Path
from typing import Any, Dict

try:
    from PIL import Image
    IMAGE_LIB_ENABLED = True
except ImportError:
    Image = None
    IMAGE_LIB_ENABLED = False

try:
    import pytesseract
    OCR_ENABLED = IMAGE_LIB_ENABLED
except ImportError:
    pytesseract = None
    OCR_ENABLED = False

class DocumentParser:
    """Parse documents and extract text, image metadata, and OCR text."""
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt'}
    
    @staticmethod
    def parse_document(file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract text plus image metadata.
        Supports: PDF, DOCX, TXT
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            text, image_count, image_details, ocr_text, images = DocumentParser._parse_pdf(file_path)
        elif file_ext == '.docx':
            text, image_count, image_details, ocr_text, images = DocumentParser._parse_docx(file_path)
        elif file_ext == '.txt':
            text, image_count, image_details, ocr_text, images = DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        output_text = text.strip() if text else ""
        image_note = ""
        if image_count:
            image_note = (
                f"\n\n[NOTE: Detected {image_count} image(s) in this document. "
                f"Details: {image_details or 'unknown'}."
            )
            if ocr_text:
                image_note += f" OCR text extracted: {ocr_text}"
            elif OCR_ENABLED:
                image_note += " No readable OCR text was found."
            else:
                image_note += " OCR is not installed, so image text may be unavailable."
            image_note += "]"

        return {
            'text': (output_text + image_note).strip(),
            'image_count': image_count,
            'image_details': image_details,
            'ocr_text': ocr_text,
            'images': images
        }
    
    @staticmethod
    def _parse_pdf(file_path: str) -> tuple[str, int, str, str, list]:
        """Extract text and image metadata from PDF"""
        try:
            text = ""
            image_count = 0
            image_details = []
            ocr_texts = []
            extracted_images = []

            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    text += page.extract_text() or ""
                    images = DocumentParser._extract_pdf_images(page)
                    if images:
                        image_count += len(images)
                        image_details.append(f"{len(images)} image(s) on page {page_num}")
                        ocr_texts.extend([image.get('ocr_text', '') for image in images if image.get('ocr_text')])
                        extracted_images.extend(images)

            details = ", ".join(image_details)
            return text, image_count, details, " | ".join(ocr_texts), extracted_images
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def _extract_pdf_images(page) -> list[Dict[str, Any]]:
        """Extract embedded images from a PDF page."""
        images = []
        try:
            resources = page.get("/Resources") or {}
            xobject = resources.get("/XObject")
            if not xobject:
                return images

            xobject = xobject.get_object()
            for obj in xobject:
                try:
                    xobj = xobject[obj]
                    if xobj.get("/Subtype") != "/Image":
                        continue
                    image_data = xobj.get_data()
                    ext = DocumentParser._detect_pdf_image_format(xobj)
                    ocr_text = DocumentParser._image_ocr(image_data)
                    images.append({
                        'name': str(obj),
                        'format': ext,
                        'data_length': len(image_data),
                        'ocr_text': ocr_text,
                        'data': image_data
                    })
                except Exception:
                    continue
        except Exception:
            pass
        return images

    @staticmethod
    def _detect_pdf_image_format(xobj) -> str:
        """Determine the image file extension from PDF object filters."""
        filters = xobj.get('/Filter')
        if isinstance(filters, list):
            filters = filters[0]
        if filters == '/DCTDecode':
            return 'jpg'
        if filters == '/JPXDecode':
            return 'jp2'
        if filters == '/FlateDecode':
            return 'png'
        return 'bin'

    @staticmethod
    def _parse_docx(file_path: str) -> tuple[str, int, str, str, list]:
        """Extract text and image metadata from DOCX"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            images = DocumentParser._extract_docx_images(doc)
            image_count = len(images)
            details = "" if image_count == 0 else f"{image_count} embedded image(s) detected"
            ocr_texts = [image['ocr_text'] for image in images if image.get('ocr_text')]
            return text, image_count, details, " | ".join(ocr_texts), images
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def _extract_docx_images(doc: Document) -> list[Dict[str, Any]]:
        """Extract embedded images from DOCX."""
        images = []
        try:
            for part in doc.part.package.parts:
                if part.content_type.startswith('image'):
                    image_data = part.blob
                    ext = part.content_type.split('/')[-1]
                    ocr_text = DocumentParser._image_ocr(image_data)
                    images.append({
                        'name': part.partname.split('/')[-1],
                        'format': ext,
                        'data_length': len(image_data),
                        'ocr_text': ocr_text,
                        'data': image_data
                    })
        except Exception:
            pass
        return images
    
    @staticmethod
    def _parse_txt(file_path: str) -> tuple[str, int, str, str, list]:
        """Extract text from TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as txt_file:
                text = txt_file.read()
            return text.strip() if text else "", 0, "", "", []
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as txt_file:
                text = txt_file.read()
            return text.strip() if text else "", 0, "", "", []
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")
    
    @staticmethod
    def _image_ocr(image_bytes: bytes) -> str:
        """Run OCR on an image if pytesseract and PIL are available."""
        if not OCR_ENABLED or not pytesseract or not IMAGE_LIB_ENABLED or not Image:
            return ""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            ocr_text = pytesseract.image_to_string(image)
            return ocr_text.strip()
        except Exception:
            return ""
    
    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """Get file information"""
        if not os.path.exists(file_path):
            return {}
        
        file_size = os.path.getsize(file_path)
        file_name = os.path.basename(file_path)
        file_ext = Path(file_path).suffix.lower()
        
        return {
            'name': file_name,
            'extension': file_ext,
            'size_bytes': file_size,
            'size_mb': round(file_size / (1024 * 1024), 2),
            'exists': True
        }
    
    @staticmethod
    def validate_file(file_path: str, max_size: int = 10 * 1024 * 1024) -> tuple:
        """
        Validate file
        Returns: (is_valid, error_message)
        """
        if not os.path.exists(file_path):
            return False, "File not found"
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > max_size:
            return False, f"File too large (max {max_size / (1024*1024)}MB)"
        
        # Check file extension
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in DocumentParser.SUPPORTED_FORMATS:
            return False, f"Unsupported format: {file_ext}. Supported: {DocumentParser.SUPPORTED_FORMATS}"
        
        return True, "OK"
    
    @staticmethod
    def cleanup_file(file_path: str) -> bool:
        """Delete temporary file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
            return True
        except Exception as e:
            print(f"❌ Error cleaning up file: {str(e)}")
            return False
